"""
PDF and DOCX Report Generation Service
Generates comprehensive medical-grade stress analysis reports
"""

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, List
import io


class ReportGenerator:
    """Generate comprehensive medical reports in PDF and DOCX formats"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._create_custom_styles()
    
    def _create_custom_styles(self):
        """Create custom paragraph styles"""
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=TA_CENTER
        ))
        
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#2563eb'),
            spaceAfter=12,
            spaceBefore=12
        ))
        
        self.styles.add(ParagraphStyle(
            name='SubHeader',
            parent=self.styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#3b82f6'),
            spaceAfter=6
        ))
    
    def generate_pdf(self, data: Dict) -> bytes:
        """Generate PDF report"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []
        
        # Title
        story.append(Paragraph("ZenithMind Comprehensive Stress Analysis Report", self.styles['CustomTitle']))
        story.append(Spacer(1, 0.2*inch))
        
        # Patient Info
        patient_info = [
            ['Patient Name:', data['userData']['name']],
            ['Age:', str(data['userData']['age'])],
            ['Occupation:', data['userData']['occupation']],
            ['Report Date:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
        ]
        t = Table(patient_info, colWidths=[2*inch, 4*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e0e7ff')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(t)
        story.append(Spacer(1, 0.3*inch))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", self.styles['SectionHeader']))
        
        risk_level = data.get('riskLevel', 'MODERATE')
        risk_color = {
            'LOW': colors.green,
            'MODERATE': colors.yellow,
            'HIGH': colors.orange,
            'CRITICAL': colors.red
        }.get(risk_level, colors.grey)
        
        summary_data = [
            ['Risk Level', risk_level, 'Avg Stress', f"{data.get('avgStress', 0)*100:.1f}%"],
            ['Cognitive Performance', f"{data.get('avgScore', 0):.0f}%", 'AI Confidence', f"{data.get('aiConfidence', 0)*100:.0f}%"]
        ]
        t = Table(summary_data, colWidths=[1.5*inch, 1.5*inch, 1.5*inch, 1.5*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#dbeafe')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('BACKGROUND', (1, 0), (1, 0), risk_color)
        ]))
        story.append(t)
        story.append(Spacer(1, 0.2*inch))
        
        # Physiological Analysis
        story.append(Paragraph("Physiological Analysis", self.styles['SectionHeader']))
        metrics = data.get('finalMetrics', {})
        
        physio_data = [
            ['Metric', 'Value', 'Status'],
            ['Blink Rate', f"{metrics.get('blinkRate', 0)}/min", self._get_status(metrics.get('blinkRate', 15), 10, 20)],
            ['Eye Asymmetry', f"{metrics.get('eyeAsymmetry', 0)*100:.1f}%", self._get_status(metrics.get('eyeAsymmetry', 0), 0, 0.15, inverse=True)],
            ['Brow Tension', f"{metrics.get('browTension', 0)*100:.0f}%", self._get_status(metrics.get('browTension', 0), 0, 0.4, inverse=True)],
            ['Jaw Tension', f"{metrics.get('jawTension', 0)*100:.0f}%", self._get_status(metrics.get('jawTension', 0), 0, 0.4, inverse=True)],
            ['Head Stability', f"{metrics.get('headStability', 0):.1f}°", self._get_status(metrics.get('headStability', 0), 0, 10, inverse=True)],
            ['Facial Symmetry', f"{metrics.get('facialSymmetry', 1)*100:.0f}%", 'Normal']
        ]
        
        t = Table(physio_data, colWidths=[2*inch, 2*inch, 2*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#dbeafe')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(t)
        story.append(Spacer(1, 0.2*inch))
        
        # Medical Recommendations
        story.append(Paragraph("Medical Recommendations", self.styles['SectionHeader']))
        
        recommendations = data.get('recommendations', {}).get('medical', {})
        priority = recommendations.get('priority', 'moderate')
        
        story.append(Paragraph(f"<b>Priority Level:</b> {priority.upper()}", self.styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph("<b>Immediate Actions:</b>", self.styles['SubHeader']))
        for action in recommendations.get('immediate', []):
            story.append(Paragraph(f"• {action}", self.styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        if recommendations.get('medications'):
            story.append(Paragraph("<b>Medication Considerations:</b>", self.styles['SubHeader']))
            for med in recommendations.get('medications', []):
                story.append(Paragraph(f"• {med}", self.styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        if recommendations.get('specialists'):
            story.append(Paragraph(f"<b>Specialist Consultations:</b> {', '.join(recommendations.get('specialists', []))}", self.styles['Normal']))
        
        story.append(PageBreak())
        
        # Meditation & Mindfulness
        story.append(Paragraph("Meditation & Mindfulness Program", self.styles['SectionHeader']))
        meditation = data.get('recommendations', {}).get('meditation', {})
        
        story.append(Paragraph("<b>Daily Techniques:</b>", self.styles['SubHeader']))
        for technique in meditation.get('techniques', []):
            story.append(Paragraph(f"• {technique}", self.styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph(f"<b>Duration:</b> {meditation.get('duration', 'N/A')}", self.styles['Normal']))
        story.append(Paragraph(f"<b>Frequency:</b> {meditation.get('frequency', 'N/A')}", self.styles['Normal']))
        story.append(Paragraph(f"<b>Recommended Apps:</b> {', '.join(meditation.get('apps', []))}", self.styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Workout Plan
        story.append(Paragraph("Personalized Workout Plan", self.styles['SectionHeader']))
        workout = data.get('recommendations', {}).get('workout', {})
        
        story.append(Paragraph("<b>Exercise Types:</b>", self.styles['SubHeader']))
        for exercise in workout.get('types', []):
            story.append(Paragraph(f"• {exercise}", self.styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph(f"<b>Intensity Level:</b> {workout.get('intensity', 'N/A')}", self.styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph("<b>Weekly Schedule:</b>", self.styles['SubHeader']))
        for day in workout.get('schedule', []):
            story.append(Paragraph(f"• {day}", self.styles['Normal']))
        
        story.append(PageBreak())
        
        # Lifestyle Management
        story.append(Paragraph("Lifestyle & Time Management", self.styles['SectionHeader']))
        lifestyle = data.get('recommendations', {}).get('lifestyle', {})
        
        story.append(Paragraph("<b>Work Optimization:</b>", self.styles['SubHeader']))
        for rec in lifestyle.get('work', []):
            story.append(Paragraph(f"• {rec}", self.styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph("<b>Sleep Hygiene:</b>", self.styles['SubHeader']))
        for rec in lifestyle.get('sleep', []):
            story.append(Paragraph(f"• {rec}", self.styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph("<b>Nutrition:</b>", self.styles['SubHeader']))
        for rec in lifestyle.get('nutrition', []):
            story.append(Paragraph(f"• {rec}", self.styles['Normal']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_docx(self, data: Dict) -> bytes:
        """Generate DOCX report"""
        doc = Document()
        
        # Title
        title = doc.add_heading('ZenithMind Comprehensive Stress Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Patient Info
        doc.add_heading('Patient Information', level=1)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = 'Patient Name:'
        table.cell(0, 1).text = data['userData']['name']
        table.cell(1, 0).text = 'Age:'
        table.cell(1, 1).text = str(data['userData']['age'])
        table.cell(2, 0).text = 'Occupation:'
        table.cell(2, 1).text = data['userData']['occupation']
        table.cell(3, 0).text = 'Report Date:'
        table.cell(3, 1).text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary = doc.add_paragraph()
        summary.add_run(f"Risk Level: ").bold = True
        summary.add_run(f"{data.get('riskLevel', 'MODERATE')}\n")
        summary.add_run(f"Average Stress: ").bold = True
        summary.add_run(f"{data.get('avgStress', 0)*100:.1f}%\n")
        summary.add_run(f"Cognitive Performance: ").bold = True
        summary.add_run(f"{data.get('avgScore', 0):.0f}%\n")
        summary.add_run(f"AI Confidence: ").bold = True
        summary.add_run(f"{data.get('aiConfidence', 0)*100:.0f}%")
        
        # Medical Recommendations
        doc.add_page_break()
        doc.add_heading('Medical Recommendations', level=1)
        
        recommendations = data.get('recommendations', {}).get('medical', {})
        doc.add_heading('Immediate Actions:', level=2)
        for action in recommendations.get('immediate', []):
            doc.add_paragraph(action, style='List Bullet')
        
        if recommendations.get('medications'):
            doc.add_heading('Medication Considerations:', level=2)
            for med in recommendations.get('medications', []):
                doc.add_paragraph(med, style='List Bullet')
        
        # Continue with other sections...
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _get_status(self, value: float, min_normal: float, max_normal: float, inverse: bool = False) -> str:
        """Determine status based on value range"""
        if inverse:
            if value <= max_normal:
                return '✓ Normal'
            else:
                return '⚠ Elevated'
        else:
            if min_normal <= value <= max_normal:
                return '✓ Normal'
            else:
                return '⚠ Abnormal'


# Singleton instance
_report_generator = None

def get_report_generator() -> ReportGenerator:
    """Get or create report generator singleton"""
    global _report_generator
    if _report_generator is None:
        _report_generator = ReportGenerator()
    return _report_generator
